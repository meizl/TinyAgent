"""文档加载器：支持 PDF、Word(docx)、TXT"""

import os
from typing import List
from tiny_agent.rag.document import Document


class BaseLoader:
    """加载器基类"""

    def load(self, file_path: str) -> List[Document]:
        raise NotImplementedError


class TextLoader(BaseLoader):
    """TXT 文件加载器"""

    def load(self, file_path: str) -> List[Document]:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        return [Document(text=text, metadata={"source": file_path, "format": "txt"})]


class PDFLoader(BaseLoader):
    """PDF 文件加载器（使用 PyPDF2）"""

    def load(self, file_path: str) -> List[Document]:
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            raise ImportError("请安装 PyPDF2: pip install PyPDF2")

        reader = PdfReader(file_path)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        return [Document(text=text, metadata={"source": file_path, "format": "pdf"})]


class WordLoader(BaseLoader):
    """Word 文件加载器（使用 python-docx）"""

    def load(self, file_path: str) -> List[Document]:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

        doc = DocxDocument(file_path)
        text = "\n".join(p.text for p in doc.paragraphs)
        return [Document(text=text, metadata={"source": file_path, "format": "docx"})]


# 格式 → 加载器映射
_LOADER_MAP = {
    ".txt": TextLoader,
    ".pdf": PDFLoader,
    ".docx": WordLoader,
    ".doc": WordLoader,
}


def load_file(file_path: str) -> List[Document]:
    """自动根据扩展名选择加载器"""
    ext = os.path.splitext(file_path)[1].lower()
    loader_cls = _LOADER_MAP.get(ext)
    if not loader_cls:
        raise ValueError(f"不支持的文件格式: {ext}，支持: {list(_LOADER_MAP.keys())}")
    return loader_cls().load(file_path)


def load_directory(dir_path: str, recursive: bool = False) -> List[Document]:
    """加载目录下所有支持的文档"""
    documents = []
    for root, _, files in os.walk(dir_path):
        for filename in files:
            ext = os.path.splitext(filename)[1].lower()
            if ext in _LOADER_MAP:
                file_path = os.path.join(root, filename)
                documents.extend(load_file(file_path))
        if not recursive:
            break
    return documents
